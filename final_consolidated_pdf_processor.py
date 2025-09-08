import os
import json
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_document(json_path, output_docx_path):
    """
    Generates a Word document from a _unified.json file.
    """
    print(f"Processing JSON file: {json_path}")
    
    # For combined_content.json, the base_dir is the folder containing it.
    base_dir = os.path.dirname(json_path)
    
    with open(json_path, 'r', encoding='utf-8') as f:
        all_content = json.load(f)
        
    document = Document()
    
    # Set some basic styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for element in all_content:
        content_type = element.get('type')
        
        if content_type == 'paragraph':
            role = element.get('role', '')
            content = element.get('content', '')
            
            # Skip empty content and standalone page numbers/footers
            if not content.strip() or role in ['pageNumber', 'pageFooter', 'pageHeader']:
                continue

            p = document.add_paragraph()
            
            # Add content run
            content_run = p.add_run(content)
            
            # Apply basic styling based on role
            if role == 'title':
                content_run.bold = True
                content_run.font.size = Pt(24)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif role == 'sectionHeading':
                content_run.bold = True
                content_run.font.size = Pt(14)
            
            # Add source text run for merged headings
            if role == 'sectionHeading' and 'source_documents' in element:
                source_docs = element.get('source_documents', [])
                source_text = f" [Source: {', '.join(source_docs)}]"
                source_run = p.add_run(source_text)
                source_run.italic = True
                source_run.font.size = Pt(9)

            # Add a little space after headings
            if role in ['title', 'sectionHeading']:
                 p.add_run('\n')


        elif content_type == 'table':
            table_data = element.get('data', [])
            if not table_data:
                continue
            
            num_rows = len(table_data)
            num_cols = len(table_data[0]) if num_rows > 0 else 0
            
            if num_rows == 0 or num_cols == 0:
                continue

            table = document.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            
            for i, row_data in enumerate(table_data):
                row_cells = table.rows[i].cells
                for j, cell_data in enumerate(row_data):
                    # Clean up checkbox markers with case-insensitive and robust logic
                    text = str(cell_data)
                    
                    # Use a case-insensitive search to see if it's selected
                    is_selected = re.search(':selected:', text, re.IGNORECASE)
                    
                    # Case-insensitively remove both tags
                    cleaned_text = re.sub(':unselected:', '', text, flags=re.IGNORECASE)
                    cleaned_text = re.sub(':selected:', '', cleaned_text, flags=re.IGNORECASE)
                    
                    # If it was selected and 'X' isn't already present, add 'X'
                    if is_selected and 'X' not in cleaned_text.upper():
                        cleaned_text = 'X' + cleaned_text
                        
                    # Final strip to clean up any whitespace
                    row_cells[j].text = cleaned_text.strip()
            
            # Add a newline after the table for spacing
            document.add_paragraph()

            # Add source tag for the table if it exists
            if 'source_documents' in element:
                p = document.add_paragraph()
                source_text = f"[Source: {', '.join(element.get('source_documents', []))}]"
                source_run = p.add_run(source_text)
                source_run.font.size = Pt(9)
                source_run.italic = True
            else:
                document.add_paragraph() # Add space after table

        elif content_type == 'image':
            image_path = element.get('path')
            if not image_path:
                continue
            
            # Construct the full path to the image
            full_image_path = os.path.join(base_dir, image_path)
            
            if os.path.exists(full_image_path):
                print(f"  - Adding image: {full_image_path}")
                # Use original width if available, otherwise default to 4.5 inches
                width_pts = element.get('width_pts')
                if width_pts:
                    document.add_picture(full_image_path, width=Pt(width_pts))
                else:
                    document.add_picture(full_image_path, width=Inches(4.5))
                
                # Add source tag for the image if it exists
                if 'source_documents' in element:
                    p = document.add_paragraph()
                    source_text = f"[Source: {', '.join(element.get('source_documents', []))}]"
                    source_run = p.add_run(source_text)
                    source_run.font.size = Pt(9)
                    source_run.italic = True
                else:
                    document.add_paragraph() # Add space after image
            else:
                print(f"  - Image not found: {full_image_path}")
                document.add_paragraph(f"[Image not found: {image_path}]", style='Caption')
            
            # Add a newline after the image for spacing
            document.add_paragraph()

    document.save(output_docx_path)
    print(f"Successfully created Word document: {output_docx_path}")


def main():
    """
    Runs the full pipeline:
    1. Extracts content from all PDFs.
    2. Combines the extracted content into a single file.
    3. Creates a final Word document from the combined content.
    """
    # --- Step 1: Run the extraction script ---
    print("--- Running Extraction Script (ext.py) ---")
    import ext
    ext.main()
    print("--- Extraction Script Finished ---\n")

    # --- Step 2: Run the combination script ---
    print("--- Running Combination Script (combine_docs.py) ---")
    import combine_docs
    combine_docs.main()
    print("--- Combination Script Finished ---\n")
    
    # --- Step 3: Create the final Word document ---
    print("--- Creating Final Combined Word Document ---")
    base_output_folder = r"C:\Users\XJ533JH\Downloads\doc_ext\extracted_content"
    combined_json_path = os.path.join(base_output_folder, "combined_content.json")
    output_docx_path = r"C:\Users\XJ533JH\Downloads\doc_ext\Final_Combined_Document.docx"
    
    if os.path.exists(combined_json_path):
        create_word_document(combined_json_path, output_docx_path)
    else:
        print(f"Error: Combined content file not found at {combined_json_path}")

if __name__ == "__main__":
    main()
